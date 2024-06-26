import streamlit as st
import assemblyai as aai
import requests
import os
import tempfile
import re
import json
import shutil
import base64
import time
from streamlit_quill import st_quill
from streamlit_javascript import st_javascript
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# AssemblyAI API key
aai.settings.api_key = "b5b3e501c6284d9dad9190aca9a4bc15"

# Create a directory for temporary files
TEMP_DIR = os.path.join(os.getcwd(), "temp_files")
os.makedirs(TEMP_DIR, exist_ok=True)

def upload_file_to_assemblyai(file_path):
    headers = {'authorization': aai.settings.api_key}
    with open(file_path, 'rb') as file:
        response = requests.post('https://api.assemblyai.com/v2/upload',
                                 headers=headers,
                                 data=file)
    return response.json()['upload_url']

def transcribe_file(file_url, word_boost, speakers_expected):
    config = aai.TranscriptionConfig(
        word_boost=word_boost,
        boost_param="high",
        speaker_labels=True,
        disfluencies=True,
        speakers_expected=speakers_expected,
        speech_model=aai.SpeechModel.best
    )
    transcriber = aai.Transcriber()
    transcript = transcriber.transcribe(file_url, config=config)
    return transcript

def save_plain_transcript(transcript):
    return json.dumps([{
        "speaker": utt.speaker,
        "text": utt.text,
        "start": utt.start,
        "end": utt.end
    } for utt in transcript.utterances])

def handle_upload_and_transcription(file, speakers_expected):
    file_id = f"{file.name}_{file.size}"
    transcript_cache_path = os.path.join(TEMP_DIR, f"transcript_cache_{file_id}.json")

    if os.path.exists(transcript_cache_path):
        st.write("Found existing transcript. Loading...")
        with open(transcript_cache_path, 'r') as cache_file:
            cached_data = json.load(cache_file)
        return cached_data['temp_file_path'], cached_data['plain_transcript']

    temp_file_path = os.path.join(TEMP_DIR, file.name)
    with open(temp_file_path, "wb") as temp_file:
        temp_file.write(file.getbuffer())

    st.write("Uploading the file to AssemblyAI...")
    file_url = upload_file_to_assemblyai(temp_file_path)
    st.write("File uploaded successfully!")
    
    st.write("Transcribing the file...")
    transcript = transcribe_file(file_url, word_boost, speakers_expected)
    st.write("Transcription completed!")
    
    plain_transcript = save_plain_transcript(transcript)

    with open(transcript_cache_path, 'w') as cache_file:
        json.dump({
            'temp_file_path': temp_file_path,
            'plain_transcript': plain_transcript
        }, cache_file)
    
    return temp_file_path, plain_transcript

def get_media_data_url(file_path):
    with open(file_path, "rb") as media_file:
        media_bytes = media_file.read()
    media_base64 = base64.b64encode(media_bytes).decode()
    file_extension = os.path.splitext(file_path)[1].lower()
    mime_type = "video/mp4" if file_extension in ['.mp4', '.avi', '.mov', '.mpeg'] else "audio/mpeg"
    return f"data:{mime_type};base64,{media_base64}"

def custom_media_player(media_file_path, transcript_data):
    media_data_url = get_media_data_url(media_file_path)
    file_extension = os.path.splitext(media_file_path)[1].lower()
    is_video = file_extension in ['.mp4', '.avi', '.mov', '.mpeg']
    
    custom_html = f"""
    <div id="media-container">
        <{'video' if is_video else 'audio'} id="media-player" controls style="width: 100%;">
            <source src="{media_data_url}" type="{'video/mp4' if is_video else 'audio/mpeg'}">
            Your browser does not support the {'video' if is_video else 'audio'} element.
        </{'video' if is_video else 'audio'}>
        <div id="transcript-display"></div>
        <div id="controls">
            <button id="slow-down">Slow Down (-0.25x)</button>
            <button id="speed-up">Speed Up (+0.25x)</button>
        </div>
    </div>
    <script>
        const mediaPlayer = document.getElementById('media-player');
        const transcriptDisplay = document.getElementById('transcript-display');
        const transcript = {transcript_data};
        const slowDownBtn = document.getElementById('slow-down');
        const speedUpBtn = document.getElementById('speed-up');

        function updateTranscript() {{
            const currentTime = mediaPlayer.currentTime;
            const currentUtterance = transcript.find(utt => currentTime >= utt.start && currentTime <= utt.end);
            
            if (currentUtterance) {{
                transcriptDisplay.innerHTML = `<strong>${{currentUtterance.speaker}}:</strong> ${{currentUtterance.text}}`;
            }}
        }}

        mediaPlayer.addEventListener('timeupdate', updateTranscript);

        function changePlaybackRate(delta) {{
            mediaPlayer.playbackRate = Math.max(0.25, Math.min(2, mediaPlayer.playbackRate + delta));
        }}

        slowDownBtn.addEventListener('click', () => changePlaybackRate(-0.25));
        speedUpBtn.addEventListener('click', () => changePlaybackRate(0.25));

        document.addEventListener('keydown', (e) => {{
            if (e.key === 'ArrowLeft') {{
                mediaPlayer.currentTime -= 5;
            }} else if (e.key === 'ArrowRight') {{
                mediaPlayer.currentTime += 5;
            }} else if (e.key === '-') {{
                changePlaybackRate(-0.25);
            }} else if (e.key === '=') {{
                changePlaybackRate(0.25);
            }}
        }});
    </script>
    <style>
        #transcript-display {{
            margin-top: 10px;
            padding: 10px;
            background-color: #f0f0f0;
            border-radius: 5px;
        }}
        #controls {{
            margin-top: 10px;
        }}
        #controls button {{
            margin-right: 10px;
        }}
    </style>
    """
    st.components.v1.html(custom_html, height=450 if is_video else 200)

def format_transcript_to_ascii(input_text):
    lines_per_page = 25
    max_line_length = 67  # Changed from 65 to 67
    formatted_lines = []
    line_number = 1
    page_number = 1

    def format_line(number, content, is_speaker=False):
        if is_speaker:
            return f"{number:2d}             {content}"
        else:
            return f"{number:2d}   {content}"

    def add_page_number():
        nonlocal page_number
        formatted_lines.append(f"{page_number:04d}")
        page_number += 1
        return 1  # Reset line number to 1

    def split_line(line, max_length):
        if len(line) <= max_length:
            return line, ""
        split_index = line.rfind(' ', 0, max_length + 1)
        if split_index == -1:
            split_index = max_length
        return line[:split_index].rstrip(), line[split_index:].lstrip()

    line_number = add_page_number()

    for line in input_text.split('\n'):
        # Remove HTML tags
        line = re.sub('<[^<]+?>', '', line)
        if ':' in line:
            speaker, text = line.split(':', 1)
            speaker = speaker.strip().upper() + ":"
            text = text.strip()

            # Add speaker name with correct spacing
            first_line, remaining = split_line(f"{speaker}     {text}", max_line_length - 13)
            formatted_lines.append(format_line(line_number, first_line, True))
            line_number += 1
            if line_number > lines_per_page:
                formatted_lines.append("\f")
                line_number = add_page_number()

            # Continue with the rest of the text
            current_line = remaining
            while current_line:
                line_content, current_line = split_line(current_line, max_line_length - 3)
                formatted_lines.append(format_line(line_number, line_content))
                line_number += 1
                if line_number > lines_per_page:
                    formatted_lines.append("\f")
                    line_number = add_page_number()

    return "\n".join(formatted_lines)

def export_to_docx(formatted_transcript):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(12)

    for line in formatted_transcript.split('\n'):
        if line.strip().isdigit():  # Page number
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line.strip()).bold = True
        elif line.strip() == '\f':  # Page break
            doc.add_page_break()
        else:
            # Remove or replace any invalid XML characters
            clean_line = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', line)
            doc.add_paragraph(clean_line)

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_docx.name)
    return temp_docx.name

def clear_cache():
    if 'temp_file_path' in st.session_state:
        del st.session_state['temp_file_path']
    if 'initial_transcript' in st.session_state:
        del st.session_state['initial_transcript']
    if 'edited_transcript' in st.session_state:
        del st.session_state['edited_transcript']
    
    # Clear the cache file
    if 'uploaded_file_name' in st.session_state and 'uploaded_file_size' in st.session_state:
        file_id = f"{st.session_state.uploaded_file_name}_{st.session_state.uploaded_file_size}"
        transcript_cache_path = os.path.join(TEMP_DIR, f"transcript_cache_{file_id}.json")
        if os.path.exists(transcript_cache_path):
            os.remove(transcript_cache_path)

# Streamlit Interface
st.title("Enhanced Transcription and Formatting Service")

# Add custom vocabulary input
st.subheader("Custom Vocabulary")
custom_vocab_input = st.text_area("Enter custom vocabulary (one word or phrase per line):", 
                                  help="Add domain-specific terms or names to improve transcription accuracy.")
word_boost = [word.strip() for word in custom_vocab_input.split('\n') if word.strip()]

# Add expected number of speakers input
speakers_expected = st.number_input("Expected number of speakers:", min_value=1, max_value=10, value=2, step=1)

uploaded_file = st.file_uploader("Upload an audio or video file", type=["mp3", "m4a", "wav", "mp4", "mpeg", "avi", "mov"])

# Add a button to clear the cache
if st.button("Clear Cached Transcript"):
    clear_cache()
    st.success("Cached transcript cleared. You can now generate a new transcript.")

if uploaded_file is not None:
    # Store file details in session state
    st.session_state.uploaded_file_name = uploaded_file.name
    st.session_state.uploaded_file_size = uploaded_file.size

    # Check if we need to generate a new transcription
    generate_new_transcript = False
    if 'temp_file_path' not in st.session_state or 'initial_transcript' not in st.session_state:
        generate_new_transcript = True

    if generate_new_transcript:
        with st.spinner('Processing...'):
            temp_file_path, initial_transcript = handle_upload_and_transcription(uploaded_file, speakers_expected)
            st.session_state.temp_file_path = temp_file_path
            st.session_state.initial_transcript = initial_transcript
        st.success("Transcription completed!")
    else:
        temp_file_path = st.session_state.temp_file_path
        initial_transcript = st.session_state.initial_transcript

    # Parse the initial transcript
    transcript_data = json.loads(initial_transcript)

    # Custom Media player with synchronized transcript
    st.subheader("Media Player with Synchronized Transcript")
    custom_media_player(temp_file_path, initial_transcript)

    st.info("Media Player Controls: Use '-' to slow down and '=' to speed up playback. Use left and right arrow keys to skip backward and forward.")

    # Editable transcript with rich text editor and auto-save
    st.subheader("Edit Transcript")
    if 'edited_transcript' not in st.session_state:
        st.session_state.edited_transcript = "\n".join([f"{utt['speaker']}: {utt['text']}" for utt in transcript_data])

    # Rich text editor with auto-save
    content = st_quill(
        value=st.session_state.edited_transcript,
        placeholder="Edit the transcript here...",
        key="quill",
    )

    # Auto-save functionality
    if content:
        st.session_state.edited_transcript = content
        st.success("Transcript auto-saved!")

    # Find and Replace
    st.subheader("Find and Replace")
    find_text = st.text_input("Find:")
    replace_text = st.text_input("Replace with:")

    if st.button("Apply Find and Replace"):
        st.session_state.edited_transcript = re.sub(re.escape(find_text), replace_text, st.session_state.edited_transcript, flags=re.IGNORECASE)
        st.experimental_rerun()

    if st.button("Format Transcript"):
        formatted_transcript = format_transcript_to_ascii(st.session_state.edited_transcript)

        st.subheader("Formatted Transcript")
        st.text_area("Formatted ASCII Transcript:", value=formatted_transcript, height=600)

        col1, col2 = st.columns(2)

        with col1:
            st.download_button(
                label="Download Formatted Transcript (TXT)",
                data=formatted_transcript,
                file_name="formatted_transcript.txt",
                mime="text/plain"
            )

        with col2:
            docx_path = export_to_docx(formatted_transcript)
            with open(docx_path, "rb") as docx_file:
                st.download_button(
                    label="Download Formatted Transcript (DOCX)",
                    data=docx_file,
                    file_name="formatted_transcript.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            os.unlink(docx_path)


# Cleanup temporary files
def cleanup_temp_files():
    shutil.rmtree(TEMP_DIR, ignore_errors=True)

# Register the cleanup function to run when the script exits
import atexit
atexit.register(cleanup_temp_files)