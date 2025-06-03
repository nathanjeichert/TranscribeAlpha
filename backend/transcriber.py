import os
import io
import json
import time
import tempfile
import logging
from typing import List, Optional

from google import genai
from google.genai import types
from google.api_core import exceptions as google_exceptions
from docx import Document
from docx.shared import Inches, Pt
import ffmpeg
from pydub import AudioSegment
from pydantic import BaseModel, ValidationError

MODEL_NAME = "gemini-2.5-pro-exp-03-25"
SUPPORTED_VIDEO_TYPES = ["mp4", "mov", "avi", "mkv"]
SUPPORTED_AUDIO_TYPES = ["mp3", "wav", "m4a", "flac", "ogg", "aac", "aiff"]

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    raise RuntimeError("GEMINI_API_KEY environment variable not set")

client = genai.Client(api_key=API_KEY)

class TranscriptTurn(BaseModel):
    speaker: str
    text: str

def convert_video_to_audio(input_path: str, output_path: str, format: str = "mp3") -> Optional[str]:
    try:
        ffmpeg.input(input_path).output(output_path, format=format, acodec='libmp3lame').overwrite_output().run(quiet=True)
        return output_path
    except ffmpeg.Error as e:
        logger.error("ffmpeg error: %s", e)
        return None


def get_audio_mime_type(ext: str) -> Optional[str]:
    mime_map = {
        "mp3": "audio/mp3",
        "wav": "audio/wav",
        "aiff": "audio/aiff",
        "aac": "audio/aac",
        "ogg": "audio/ogg",
        "flac": "audio/flac",
    }
    return mime_map.get(ext.lower())


def upload_to_gemini(file_path: str) -> Optional[types.File]:
    try:
        gemini_file = client.files.upload(file=file_path)
        file_state = "PROCESSING"
        retries = 15
        sleep_time = 8
        max_sleep = 45
        while file_state == "PROCESSING" and retries > 0:
            time.sleep(sleep_time)
            file_info = client.files.get(name=gemini_file.name)
            file_state = file_info.state.name
            retries -= 1
            sleep_time = min(sleep_time * 1.5, max_sleep)
        if file_state != "ACTIVE":
            try:
                client.files.delete(name=gemini_file.name)
            except Exception:
                pass
            return None
        return gemini_file
    except Exception as e:
        logger.error("upload failed: %s", e)
        return None


def generate_transcript(gemini_file: types.File, speaker_name_list: Optional[List[str]] = None) -> Optional[List[TranscriptTurn]]:
    safety_settings = [
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HARASSMENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_CIVIC_INTEGRITY, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    ]

    if speaker_name_list:
        speaker_prompt_part = f"The speakers are identified as: {', '.join(speaker_name_list)}."
        num_speakers_part = f"There are {len(speaker_name_list)} speakers."
    else:
        speaker_prompt_part = "Speaker identifiers are not provided; use generic identifiers like SPEAKER 1, SPEAKER 2, etc., IN ALL CAPS."
        num_speakers_part = "Determine the number of speakers from the audio."

    prompt = (
        f"Generate a transcript of the speech. {num_speakers_part} {speaker_prompt_part} "
        "Structure the output STRICTLY as a JSON list of objects. "
        "Each object represents a continuous block of speech from a single speaker and MUST contain BOTH a 'speaker' field "
        "and a 'text' field containing ALL consecutive speech from that speaker before the speaker changes."
    )

    contents = [prompt, gemini_file]
    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=contents,
            config=types.GenerateContentConfig(
                safety_settings=safety_settings,
                response_mime_type="application/json",
                response_schema=list[TranscriptTurn],
            ),
        )
        transcript_data = json.loads(response.text)
        validated_turns = []
        for turn_data in transcript_data:
            if 'speaker' not in turn_data:
                continue
            if 'text' not in turn_data:
                turn_data['text'] = ""
            try:
                validated_turns.append(TranscriptTurn(**turn_data))
            except ValidationError:
                continue
        return validated_turns
    except Exception as e:
        logger.error("generate_transcript failed: %s", e)
        return None


def replace_placeholder_text(element, placeholder: str, replacement: str) -> None:
    if hasattr(element, 'paragraphs'):
        for p in element.paragraphs:
            replace_placeholder_text(p, placeholder, replacement)
    if hasattr(element, 'runs'):
        if placeholder in element.text:
            inline = element.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    text = inline[i].text.replace(placeholder, replacement)
                    inline[i].text = text
    if hasattr(element, 'tables'):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholder_text(cell, placeholder, replacement)


def create_docx(title_data: dict, transcript_turns: List[TranscriptTurn]) -> bytes:
    doc = Document("transcript_template.docx")
    for key, value in title_data.items():
        placeholder = f"{{{{{key}}}}}"
        replace_placeholder_text(doc, placeholder, str(value) if value else "")

    body_placeholder = "{{TRANSCRIPT_BODY}}"
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if body_placeholder in p.text:
            placeholder_paragraph = p
            break

    if placeholder_paragraph:
        p_element = placeholder_paragraph._element
        p_element.getparent().remove(p_element)
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
            speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
            speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def process_transcription(file_bytes: bytes, filename: str, speaker_names: Optional[List[str]], title_data: dict):
    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = os.path.join(temp_dir, filename)
        with open(input_path, "wb") as f:
            f.write(file_bytes)

        ext = filename.split('.')[-1].lower()
        audio_path = None
        if ext in SUPPORTED_VIDEO_TYPES:
            output_audio_filename = f"{os.path.splitext(filename)[0]}.mp3"
            output_path = os.path.join(temp_dir, output_audio_filename)
            audio_path = convert_video_to_audio(input_path, output_path)
            ext = "mp3"
        elif ext in SUPPORTED_AUDIO_TYPES:
            audio_path = input_path
        else:
            raise ValueError("Unsupported file type")

        mime_type = get_audio_mime_type(ext)
        audio_segment = AudioSegment.from_file(audio_path)
        duration_seconds = len(audio_segment) / 1000
        hours, rem = divmod(duration_seconds, 3600)
        minutes, seconds = divmod(rem, 60)
        file_duration_str = "{:0>2}:{:0>2}:{:0>2}".format(int(hours), int(minutes), int(round(seconds)))
        title_data["FILE_DURATION"] = file_duration_str

        gemini_file = upload_to_gemini(audio_path)
        if not gemini_file:
            raise RuntimeError("Failed to upload file to Gemini")
        turns = generate_transcript(gemini_file, speaker_names)
        client.files.delete(name=gemini_file.name)
        if not turns:
            raise RuntimeError("Failed to generate transcript")
        docx_bytes = create_docx(title_data, turns)
        return turns, docx_bytes
